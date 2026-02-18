from datetime import datetime
from io import BytesIO
from typing import List

from googleapiclient.http import MediaIoBaseDownload

from app.config import settings
from ingestion.gdrive_client import get_drive_service, list_files_in_folders
from models.document import Document


def _download_file_content(service, file_id: str) -> bytes:
    """Download raw file bytes from Drive."""
    request = service.files().get_media(fileId=file_id)
    fh = BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return fh.getvalue()


def _export_google_doc_as_text(service, file_id: str) -> str:
    """Export a Google Doc to plain text."""
    request = service.files().export_media(fileId=file_id, mimeType="text/plain")
    fh = BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return fh.getvalue().decode("utf-8", errors="ignore")


def ingest_gdrive() -> List[Document]:
    """Ingest configured Google Drive folders into Document objects.

    This initial version focuses on Google Docs and simple text-like files.
    """
    import streamlit as st
    print("Starting Google Drive Ingestion...")
    st.toast("📂 Starting Google Drive ingestion...")
    
    root_ids = [fid.strip() for fid in settings.gdrive_root_folder_ids if fid.strip()]
    if not root_ids:
        msg = "Warning: No Google Drive root folder IDs configured in .env (GDRIVE_ROOT_FOLDER_IDS). Skipping Google Drive."
        print(msg)
        st.warning(msg)
        return []
        
    print(f"configured folder IDs: {root_ids}")

    try:
        service = get_drive_service()
        files = list_files_in_folders(service, root_ids)
        st.toast(f"📂 Google Drive: Found {len(files)} files found.")
    except Exception as e:
        msg = f"Skipping Google Drive ingestion: {e}"
        print(msg)
        st.error(f"Google Drive Failed: {e}")
        return []

    documents: List[Document] = []

    for file in files:
        file_id = file["id"]
        name = file.get("name", "Untitled")
        mime_type = file.get("mimeType", "")

        text_content = ""

        if mime_type == "application/vnd.google-apps.document":
            text_content = _export_google_doc_as_text(service, file_id)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": # DOCX
            raw = _download_file_content(service, file_id)
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(raw))
            text_content = "\n".join([p.text for p in doc.paragraphs])
        elif mime_type == "application/pdf": # PDF
            try:
                raw = _download_file_content(service, file_id)
                import pdfplumber
                from io import BytesIO
                with pdfplumber.open(BytesIO(raw)) as pdf:
                    text_content = "\n".join([(page.extract_text() or "") for page in pdf.pages])
                
                print(f"📄 DEBUG: Extracted {len(text_content)} chars from '{name}'. Preview: {text_content[:50].replace('\n', ' ')}...")
                
                if not text_content.strip():
                    st.warning(f"⚠️ Warning: PDF '{name}' seems empty or image-based (no text extracted).")
            except Exception as e:
                msg = f"Failed to read PDF '{name}': {e}. Skipping."
                print(msg)
                st.error(msg)
                continue
        elif mime_type in ("text/plain", "text/markdown"):
            raw = _download_file_content(service, file_id)
            text_content = raw.decode("utf-8", errors="ignore")
        elif mime_type == "application/vnd.google-apps.folder":
            print(f"Skipping subfolder '{name}' (recursive search not enabled yet). Add this folder ID directly to .env if needed.")
            continue
        else:
            print(f"Skipping unsupported file: {name} ({mime_type})")
            continue

        metadata = (
            service.files()
            .get(fileId=file_id, fields="id, name, mimeType, createdTime, modifiedTime, webViewLink, parents")
            .execute()
        )

        created_at = metadata.get("createdTime")
        modified_at = metadata.get("modifiedTime")

        documents.append(
            Document(
                id=f"gdrive_{file_id}",
                source="gdrive",
                source_id=file_id,
                title=name,
                text=text_content,
                created_at=datetime.fromisoformat(created_at.replace("Z", "+00:00"))
                if created_at
                else None,
                updated_at=datetime.fromisoformat(modified_at.replace("Z", "+00:00"))
                if modified_at
                else None,
                url=metadata.get("webViewLink"),
                path="",
                metadata_raw=metadata,
            )
        )

    return documents

